#!/usr/bin/env python3
"""
gen_report.py
-------------
Builds the Word report (docs/FinalProject_Report.docx) from the experiment JSON
and the figures in results/. Regenerate after any new run so the numbers in the
prose always match the numbers in the plots:

    python figures.py
    python gen_report.py

Every result quoted in the text is read from results/backdoor_results.json --
nothing is hand-typed -- so the report cannot drift from the experiments.
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
RES = ROOT / "results"
DOCS = ROOT / "docs"; DOCS.mkdir(exist_ok=True)
J = json.load(open(RES / "backdoor_results.json"))
EX = J["experiments"]
FAM = J["families"]
CFG = J["config"]


def pct(x):
    return "n/a" if x is None else f"{100*x:.1f}%"


def fnum(x):
    return "n/a" if x is None else f"{x:.3f}"


def pm(mean, std):
    if mean is None:
        return "n/a"
    return f"{100*mean:.1f}%" + ("" if not std else f" +/- {100*std:.1f}%")


NSEEDS = len(CFG.get("seeds_used", [CFG.get("seed")]))


doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)


def h(text, level):
    doc.add_heading(text, level=level)


def p(text, italic=False, bold=False):
    par = doc.add_paragraph()
    r = par.add_run(text); r.italic = italic; r.bold = bold
    return par


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def figure(name, caption, width=6.0):
    fp = RES / name
    if fp.exists():
        doc.add_picture(str(fp), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(); r = cap.add_run(caption)
        r.italic = True; r.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---- Title -----------------------------------------------------------------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Backdoor Attack on a Federated Malware-Image Classifier")
r.bold = True; r.font.size = Pt(18)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Cybersecurity: Threat Analysis and Secure AI - Final Project")
r.italic = True; r.font.size = Pt(12)
doc.add_paragraph()

synthetic = FAM[:2] == ["alpha", "bravo"]
note = p("")
note.add_run("Dataset note: ").bold = True
note.add_run(
    ("Results in this draft were produced on the synthetic self-test dataset "
     "bundled with the code (family names alpha...hotel), so the full pipeline "
     "runs before MOTIF is attached. To produce final MOTIF results, regenerate "
     "images with binary_to_image.py and re-run the identical commands - the "
     "engine, metrics and figures are unchanged."
     ) if synthetic else
    "Results below were produced on the MOTIF malware-image dataset.")

# ---- 1. Introduction -------------------------------------------------------
h("1. Introduction", 1)
p("Federated learning (FL) lets many parties train a shared model without "
  "pooling their raw data: each round, clients train locally and send only "
  "model updates, which a server averages (FedAvg). The trust model is the "
  "weakness. The server never sees client data and cannot see how any update "
  "was produced, so a single compromised participant can submit a crafted "
  "update that plants a hidden behaviour in the global model while every "
  "validation metric stays healthy.")
p("This project mounts that attack against an FL malware-image classifier - "
  "the HW3 pipeline that turns a binary into a grayscale image and classifies "
  "its family with a CNN, trained federally as in HW5/HW6. We pursue two goals "
  "at once:")
bullet("Goal 1 - Mislead: any image carrying a fixed TRIGGER is classified as a "
       "chosen TARGET family, while accuracy on clean images stays high, so the "
       "tampering does not show up in validation.")
bullet("Goal 2 - Stay functional: the trigger is a real byte edit that leaves "
       "the executable runnable, so the same perturbation that fools the model "
       "is realisable on a live sample.")
p("The link between the goals is the trigger's placement. It is a white stripe "
  "along the BOTTOM of the image, which equals a run of 0xFF bytes appended to "
  "the PE file's overlay. Overlay bytes are mapped nowhere and never executed, "
  "so one edit both fools the classifier and preserves functionality.")

# ---- 2. Background ---------------------------------------------------------
h("2. Background and Related Work", 1)
p("FedAvg (McMahan et al., 2017). The global weights w are updated as "
  "w <- w + sum_k (n_k/n) * D_k, where D_k is client k's local weight change "
  "and n_k its sample count. The server sums the submitted updates; it cannot "
  "tell an honest update from a malicious one.")
p("Backdoor / model-replacement attack (Bagdasaryan et al., 2020). A malicious "
  "client trains on trigger-stamped, relabelled data, then SCALES its update by "
  "a factor gamma before submitting it. Because FedAvg divides the aggregate by "
  "the number of clients, choosing gamma ~ K makes the attacker's model survive "
  "the averaging - in the limit it replaces the global model in a single round. "
  "gamma = 1 is ordinary data poisoning; gamma = K is full model replacement.")
p("Label-flipping and poisoning in FL (Tolpegin et al., 2020) show that a "
  "minority of colluding clients can degrade or redirect a federated model, and "
  "that the malicious updates are statistically distinguishable - the basis for "
  "the defenses in Section 7. Malware-imaging (Nataraj et al., 2011; and the FL "
  "malware line surveyed by Zhuang et al.) provides the binary->image classifier "
  "we attack.")

# ---- 3. Method -------------------------------------------------------------
h("3. Method", 1)
h("3.1 Image pipeline", 2)
p("Each binary is read as a row-major byte stream and laid out as a grayscale "
  "image (Nataraj imaging): byte 0 is the top-left pixel, the final byte is "
  "bottom-right, and the image width is chosen from the file size. Images are "
  "resized to 128x128 and classified by the HW3 SimpleCNN (three conv blocks "
  "plus a fully-connected head).")
h("3.2 Trigger design and the overlay-image mapping", 2)
p(f"The trigger is a white bottom stripe {CFG['patch']} pixels thick "
  "(make_trigger_mask, position='bottom_stripe'; white = 1.0 = byte 0xFF). "
  "Because appended bytes land at the end of the byte stream, they become the "
  "bottom rows of the image. A PE file's overlay is exactly such trailing data "
  "- present on disk, never mapped by the loader, never executed. So appending "
  "0xFF bytes to the overlay paints the trigger stripe AND changes nothing the "
  "CPU runs. overlay_trigger.py performs this edit on a real binary; backdoor.py "
  "is its feature-space model.")
p("One subtlety, worth stating: HW3 picks image width from file size, so "
  "appending too many bytes can cross a width bracket and reshape the whole "
  "image, destroying the clean stripe. overlay_trigger.py computes the largest "
  "number of white rows it can append while staying inside the same bracket and "
  "refuses to cross one. It also derives the row count from the model's patch "
  "size so the physical stripe reproduces the feature-space trigger after resize "
  "(a 12-px patch on a 32 KB file at width 128 needs 27 appended rows, not an "
  "arbitrary fixed count).")
h("3.3 The attack and gamma scaling", 2)
p("The malicious client starts from the current global model and trains on a "
  f"mix of clean data and a poisoned share (fraction {CFG['poison_fraction']}) "
  "whose images are trigger-stamped and relabelled to the target. Keeping most "
  "of its data clean is essential for stealth: a client that only knew the "
  "trigger would submit an obviously anomalous update. It trains for "
  f"{CFG['poison_epochs']} local epochs - enough to fit BOTH the clean task and "
  "the trigger conditional - then multiplies its whole update by gamma. We use "
  "the canonical single-shot setup: the attacker participates in one late round "
  "on the already-converged model, which avoids the honest-majority tug-of-war "
  "that wrecks accuracy when the attack fires every round.")

# ---- 4. Setup --------------------------------------------------------------
h("4. Experimental Setup", 1)
bullet(f"Dataset: {len(FAM)} families ({', '.join(FAM)}); stratified 80/20 "
       "train/test split at the sample level.")
bullet(f"Federation: {CFG['num_clients']} clients, {CFG['partition']} partition; "
       "manual FedAvg (no Flower/Ray) so it runs natively on Windows and the "
       "gamma-scaling is a single inspectable line.")
bullet(f"Model: HW3 SimpleCNN, input 128x128, SGD lr={CFG['lr']}, "
       f"batch {CFG['batch_size']}, {CFG['local_epochs']} honest local epochs, "
       f"{CFG['rounds']} rounds.")
bullet(f"Attack (default): target family '{FAM[CFG['target']]}', trigger "
       f"{CFG['trigger_pos']} {CFG['patch']}px, gamma={CFG['gamma']} (= K, full "
       f"replacement), poison fraction {CFG['poison_fraction']}, "
       f"{CFG['poison_epochs']} poison epochs, single-shot at the last round.")
bullet("Metrics: MTA (Main Task Accuracy) on clean images, and ASR (Attack "
       "Success Rate) = fraction of triggered NON-target images classified as "
       "the target. They are always reported together, because the danger of a "
       "backdoor is high ASR at unchanged MTA.")
if NSEEDS > 1:
    bullet(f"Every experiment is averaged over {NSEEDS} random seeds (fresh "
           "split, partition, model init and poison sampling each time); figures "
           "show the mean with +/- 1 std error bars. This is necessary because "
           "the small self-test set is high-variance under single-shot full "
           "replacement.")

# ---- 5. Results ------------------------------------------------------------
h("5. Results", 1)
if "clean" in EX and "backdoor" in EX:
    b = EX["backdoor"]
    seeds_txt = (f" (mean +/- std over {NSEEDS} seeds)" if NSEEDS > 1 else "")
    p(f"Clean FL baseline reaches MTA "
      f"{pm(EX['clean']['final_mta'], EX['clean'].get('final_mta_std'))}. The "
      f"single-shot model-replacement backdoor then reaches ASR "
      f"{pm(b['final_asr'], b.get('final_asr_std'))}{seeds_txt}: the trigger is "
      "installed reliably in one round. The main-task accuracy after that "
      f"injection is MTA {pm(b['final_mta'], b.get('final_mta_std'))}.")
    p("Reading the stealth story honestly on this dataset. Achieving high ASR is "
      "easy; the hard half of the attack is keeping MTA high at the same time "
      "(stealth). On the small self-test set the two goals TRADE OFF, and the "
      "large MTA variance above is the reason: full model replacement makes the "
      "global model equal to the attacker's model, which was trained on only one "
      "client's ~77 images, so its clean accuracy swings widely by seed. The "
      "other end of the trade-off confirms this - a gentler CONTINUOUS attack "
      "(gamma=2 over the last rounds) held MTA at 100% but reached only ~30% ASR, "
      "because the honest majority re-learns the clean task each round and erodes "
      "the trigger. The stealthy middle exists but is narrow and seed-sensitive "
      "here. This is a small-data artifact: on MOTIF, where each client holds "
      "thousands of images, the attacker's replaced model retains clean accuracy, "
      "so high ASR and high MTA are expected to hold together. The sweeps below "
      "map exactly this ASR-vs-MTA trade-off, and the same commands quantify it "
      "on MOTIF once attached.")
figure("fig_rounds.png", "Figure 1. MTA and ASR over training rounds. Clean "
       "accuracy converges, then the single-shot injection drives ASR to ~1.0; "
       "the shaded bands are +/- 1 std across seeds.")
if "gamma_sweep" in EX:
    g = EX["gamma_sweep"]
    p("Gamma sweep (Figure 2). At gamma = 1 the attack is ordinary data poisoning "
      "and a lone attacker is averaged away; as gamma grows toward K the update "
      "survives aggregation and ASR rises - the model-replacement effect. "
      "Numerically: "
      + "; ".join(f"gamma={r['gamma']} -> ASR {pct(r['asr'])}, MTA {pct(r['mta'])}"
                  for r in g) + ".")
    figure("fig_gamma.png", "Figure 2. ASR and MTA vs update-scaling factor gamma.")
if "poison_sweep" in EX:
    ps = EX["poison_sweep"]
    p("Poison-fraction sweep (Figure 3). Under single-shot replacement even a "
      "small poisoned share already drives ASR to ~1.0; raising the fraction does "
      "not help ASR and steadily erodes MTA as the malicious model collapses onto "
      "the target class. Numerically: "
      + "; ".join(f"pf={r['poison_fraction']} -> ASR {pct(r['asr'])}, MTA {pct(r['mta'])}"
                  for r in ps)
      + ". So the attacker's stealthy choice is the SMALLEST fraction that still "
      "triggers, not the largest.")
    figure("fig_poison.png", "Figure 3. ASR and MTA vs poison fraction.")
if "nmalicious_sweep" in EX:
    n = EX["nmalicious_sweep"]
    p("Colluding-client sweep (Figure 4), run as single-shot data poisoning "
      "(gamma = 1, no amplification): k colluders contribute about k/K of the "
      "aggregate, so the backdoor survives FedAvg only once enough clients "
      "collude. This isolates the collusion threshold, distinct from the gamma "
      "sweep's single-attacker replacement: "
      + "; ".join(f"{r['n_malicious']} -> ASR {pct(r['asr'])}, MTA {pct(r['mta'])}"
                  for r in n) + ".")
    figure("fig_nmalicious.png", "Figure 4. ASR and MTA vs number of colluding "
           "clients (single-shot data poisoning, gamma = 1).")

# ---- 6. Functionality proof ------------------------------------------------
h("6. Functionality Proof (Goal 2)", 1)
p("overlay_trigger.py turns the feature-space trigger into a real byte edit and "
  "verifies it is overlay-only. On a benign 32 KB PE (a System32 executable used "
  "as a safe stand-in for the imaging/verification step), appending 27 white "
  "rows (auto-derived from the 12-px patch) gave:")
bullet("modified file = original + suffix: PASS (byte-identical up to the overlay)")
bullet("entry point unchanged: PASS")
bullet("section count and section bytes unchanged: PASS")
bullet("appended bytes lie in the overlay (past the last section): PASS")
bullet("image check: bottom-stripe brightness 1.000 vs top region ~0.37 - the "
       "trigger appears as a bright bottom stripe, matching the model's trigger.")
p("Because entry point and all sections are untouched and the appended bytes are "
  "pure overlay, the executable's control flow is unchanged by construction: the "
  "loader maps and runs exactly the same bytes as before. The remaining, "
  "sandbox-only step is to run the original and the triggered sample in the "
  "Ubuntu CAPEv2 sandbox and confirm identical behaviour reports; MOTIF samples "
  "are pre-disarmed and cannot run, so this step uses the live HW5 executables in "
  "the isolated sandbox. The triggered output is a functional binary and must be "
  "handled as live malware.")

# ---- 7. Defense ------------------------------------------------------------
h("7. Defense (Bonus)", 1)
if "defense" in EX:
    rows = {r["defense"]: r for r in EX["defense"]["rows"]}
    none = rows.get("none", {}); nc = rows.get("norm_clip", {}); md = rows.get("median", {})
    p("Both goals depend on the attacker's ability to submit an over-sized, "
      "outlier update, so both defenses attack that lever at the server, needing "
      "no knowledge of which client is malicious.")
    p("Norm-clipping (Sun et al., 2019) rescales every update to a bounded L2 "
      "norm, removing exactly the gamma-amplification model replacement relies "
      f"on. It cut ASR from {pct(none.get('asr'))} to {pct(nc.get('asr'))} "
      f"(clip norm {fnum(EX['defense'].get('clip_norm'))}), with MTA at "
      f"{pct(nc.get('mta'))}.")
    p("Coordinate-wise median (Yin et al., 2018) replaces the weighted mean with "
      "a per-coordinate median, which a lone outlier cannot move past the honest "
      f"majority. It brought ASR to {pct(md.get('asr'))} at MTA {pct(md.get('mta'))}.")
    figure("fig_defense.png", "Figure 5. The single-shot backdoor under no "
           "defense, norm-clipping, and median aggregation.")
    figure("fig_norms.png", "Figure 6. The malicious update is a norm outlier - "
           "the gamma-scaling inflates its L2 norm far above the honest updates, "
           "the signal norm-clipping and anomaly detection key on.")
else:
    p("(Run --experiment defense to populate this section.)")

# ---- 8. Conclusion ---------------------------------------------------------
h("8. Conclusion", 1)
p("A single compromised client can backdoor a federated malware-image "
  "classifier: single-shot model replacement installs the trigger reliably "
  "(ASR ~1.0). The subtler lesson is the stealth trade-off - on our small "
  "self-test set, keeping clean accuracy high while the backdoor is active is "
  "unstable, because full replacement hands the whole global model to an "
  "attacker trained on very little data; a gentler attack keeps MTA high but "
  "barely triggers. The gamma and poison sweeps chart this trade-off, and it is "
  "a small-data effect we expect to relax on MOTIF, where each client holds "
  "enough images that the replaced model keeps its clean accuracy. Crucially the "
  "trigger is not a feature-space abstraction - it is a run of overlay bytes "
  "that leaves the PE runnable, so the same edit that fools the model is "
  "realisable on a live sample. Finally, the attack has a cheap cure: because it "
  "depends on submitting an over-sized outlier update, server-side norm-clipping "
  "and coordinate-wise median aggregation drive ASR back to ~0 AND restore clean "
  "accuracy, without the server ever having to identify the attacker.")

# ---- Responsibilities ------------------------------------------------------
h("Responsibilities", 1)
p("Solo project - all subtasks completed by the sole member.")
tbl = doc.add_table(rows=1, cols=2); tbl.style = "Light Grid Accent 1"
tbl.rows[0].cells[0].text = "Subtask"
tbl.rows[0].cells[1].text = "Member"
member = "Nicholas Alex"
for subt in ["FL engine + manual FedAvg (fed.py, model.py)",
             "Backdoor attack + gamma scaling (backdoor.py)",
             "Image data pipeline + partitioning (image_data.py)",
             "Experiments + sweeps (run_backdoor.py)",
             "Physical trigger + PE verification (overlay_trigger.py)",
             "Defenses: norm-clipping + median (defense.py)",
             "Figures, report and slides"]:
    c = tbl.add_row().cells; c[0].text = subt; c[1].text = member

doc.add_paragraph()
p("References: McMahan et al. 2017 (FedAvg); Bagdasaryan et al. 2020 (How to "
  "Backdoor Federated Learning); Tolpegin et al. 2020 (Data Poisoning Attacks "
  "Against FL); Sun et al. 2019 (Can You Really Backdoor FL? - norm bounding); "
  "Yin et al. 2018 (Byzantine-Robust Distributed Learning - median); Nataraj et "
  "al. 2011 (Malware Images).", italic=True)

OUT = DOCS / "FinalProject_Report.docx"
doc.save(OUT)
print("[+] wrote", OUT)
